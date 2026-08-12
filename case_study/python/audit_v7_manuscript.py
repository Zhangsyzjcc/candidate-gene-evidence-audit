#!/usr/bin/env python3
from __future__ import annotations

import csv
import re
import subprocess
import sys
import zipfile
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "manuscript/main/LRRK2_glioma_full_manuscript_v7_en_2026-08-04.md"
DOCX = MD.with_suffix(".docx")
OUT = ROOT / "results/qc/technical_tests/v7_manuscript_audit_2026-08-04.csv"


def main() -> None:
    checks: list[tuple[str, bool, str]] = []
    def check(name: str, ok: bool, detail: str = "") -> None:
        checks.append((name, bool(ok), detail))

    required_files = [
        MD, DOCX,
        ROOT / "results/tables/supplementary/Table_S6_MYC_DNA_repair_p53_axis_models_2026-08-03.csv",
        ROOT / "results/tables/supplementary/Table_S7_MYC_DNA_repair_p53_consensus_leading_edge_2026-08-03.csv",
        ROOT / "results/figures/v7/supplementary/FigS3_MYC_DNA_repair_p53_axis/FigS3_MYC_DNA_repair_p53_axis_2026-08-04.pdf",
        ROOT / "results/figures/v7/supplementary/FigS3_MYC_DNA_repair_p53_axis/FigS3_MYC_DNA_repair_p53_axis_2026-08-04.png",
        ROOT / "results/figures/v7/supplementary/FigS3_MYC_DNA_repair_p53_axis/FigS3_MYC_DNA_repair_p53_axis_2026-08-04.svg",
        ROOT / "manuscript/claim_evidence/v7_damage_axis_claims_2026-08-04.csv",
    ]
    for path in required_files:
        check(f"file_{path.name}", path.is_file() and path.stat().st_size > 0, path.relative_to(ROOT).as_posix())

    text = MD.read_text(encoding="utf-8")
    required_phrases = [
        "Sample-level MYC-DNA repair-p53/damage-response axis",
        "beta=-0.327", "BH-adjusted P=0.133", "BH-adjusted P=0.0198",
        "138 program-gene pairs", "partial sample-level replication",
        "different estimands", "Supplementary Tables S6-S7", "Supplementary Figure S3",
        "do not establish regulation by LRRK2",
    ]
    for phrase in required_phrases:
        check(f"phrase_{phrase[:24]}", phrase in text, phrase)
    check("immune_S2_retained", text.count("Supplementary Figure S2") == 1, f"count={text.count('Supplementary Figure S2')}")
    check("new_S3_referenced", text.count("Supplementary Figure S3") >= 3, f"count={text.count('Supplementary Figure S3')}")
    check("no_mechanism_claim", "LRRK2 regulates the" not in text and "LRRK2 drives" not in text, "causal phrases absent")
    refs = re.findall(r"(?m)^\d+\. ", text[text.index("# References"):text.index("# Figure legends")])
    check("reference_count_preserved", len(refs) == 21, f"references={len(refs)}")

    check("docx_zip", zipfile.is_zipfile(DOCX), "valid OOXML zip")
    doc = Document(str(DOCX))
    check("docx_one_section", len(doc.sections) == 1, f"sections={len(doc.sections)}")
    check("docx_paragraphs", len(doc.paragraphs) >= 140, f"paragraphs={len(doc.paragraphs)}")
    doc_text = "\n".join(p.text for p in doc.paragraphs)
    check("docx_required_content", all(x in doc_text for x in ("Supplementary Figure S3", "beta=-0.327", "partial sample-level replication")), "V7 additions present")
    with zipfile.ZipFile(DOCX) as z:
        xml = "".join(z.read(n).decode("utf-8", "ignore") for n in z.namelist() if n.endswith(".xml"))
        check("docx_no_comments", "word/comments.xml" not in z.namelist(), "comments absent")
        check("docx_no_tracked_changes", re.search(r"<w:(?:ins|del)(?:[ >])", xml) is None, "tracked changes absent")

    preservation = subprocess.run([sys.executable, str(ROOT / "python/audit_v6_manuscript_preservation.py")], cwd=ROOT,
                                  capture_output=True, text=True)
    check("v6_preservation", preservation.returncode == 0, preservation.stdout.strip() or preservation.stderr.strip())

    OUT.parent.mkdir(parents=True, exist_ok=True)
    with OUT.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f); w.writerow(["check", "passed", "detail"])
        w.writerows((n, str(ok).lower(), d) for n, ok, d in checks)
    failed = [n for n, ok, _ in checks if not ok]
    print(f"checks={len(checks)} failures={len(failed)}")
    for name in failed:
        print("FAIL", name)
    raise SystemExit(1 if failed else 0)


if __name__ == "__main__":
    main()
